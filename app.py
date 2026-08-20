import streamlit as st
import pandas as pd
import re
import io
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# --- SETTINGS & UI CONFIG ---
st.set_page_config(page_title="Disfluency Analyzer", layout="wide")
st.title("🗣️ Speech Disfluency Analyzer (v1.7)")

# --- HELPERS ---
def get_seconds(time_str):
    try:
        parts = [int(p) for p in time_str.split(':')]
        if len(parts) == 3: return parts[0] * 3600 + parts[1] * 60 + parts[2]
        if len(parts) == 2: return parts[0] * 60 + parts[1]
        return 0
    except: return 0

def extract_timestamp(line):
    match = re.search(r'(\d{1,2}:\d{2}:\d{2})|(\d{1,2}:\d{2})', line)
    return get_seconds(match.group(0)) if match else None

def clean_transcript_clutter(text, exclusion_list):
    # 1. Basic Cleaning (Timestamps and VTT markers)
    text = re.sub(r'\d{1,2}:\d{2}:\d{2}\.\d{3} --> \d{1,2}:\d{2}:\d{2}\.\d{3}', '', text)
    text = re.sub(r'\b\d{1,2}:\d{2}(:\d{2})?\b', '', text)
    text = re.sub(r'^[A-Za-z\s\d]+:', '', text, flags=re.MULTILINE)
    
    # 2. Universal Exclusions (Researcher talk, intros, countdowns)
    for phrase in exclusion_list:
        p_clean = phrase.strip()
        if p_clean:
            # Flexible regex: Escapes phrase and allows for varying punctuation/spaces
            pattern = re.escape(p_clean).replace(r'\ ', r'\s+').replace(r'\,', r'[\s,.]*')
            text = re.sub(pattern + r'[\s,.]*', '', text, flags=re.IGNORECASE)
            
    text = re.sub(r'\s+', ' ', text)
    return text.strip(",. ")

def build_reviewed_speech_docx(participant_id, observer_id, visit_date, speech_results):
    doc = Document()
    doc.add_heading("Reviewed Speech Summary", 0)
    doc.add_paragraph(f"Participant ID: {participant_id}")
    doc.add_paragraph(f"Observer ID: {observer_id}")
    doc.add_paragraph(f"Visit Date: {visit_date.strftime('%B %d, %Y')}")
    doc.add_paragraph("Researcher-approved disfluencies are highlighted in yellow.")

    for speech in speech_results:
        doc.add_heading(speech["name"], level=1)
        paragraph = doc.add_paragraph()
        text = speech["text"]
        cursor = 0
        for finding in sorted(speech["approved_findings"], key=lambda item: item["start"]):
            start, end = finding["start"], finding["end"]
            if start < cursor:
                continue
            paragraph.add_run(text[cursor:start])
            highlighted_run = paragraph.add_run(text[start:end])
            highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            cursor = end
        paragraph.add_run(text[cursor:])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def build_token_coding_rows(participant_id, observer_id, visit_date, speech_results):
    rows = []
    for speech in speech_results:
        approved = sorted(speech["approved_findings"], key=lambda item: item["start"])
        occurrence_ids = {
            (finding["start"], finding["end"]): f"S{speech['speech_number']}-D{number:03d}"
            for number, finding in enumerate(approved, start=1)
        }
        for token_position, match in enumerate(re.finditer(r"\S+", speech["text"]), start=1):
            token_start, token_end = match.start(), match.end()
            coded_finding = next(
                (
                    finding for finding in approved
                    if token_start < finding["end"] and token_end > finding["start"]
                ),
                None,
            )
            rows.append({
                "Participant_ID": participant_id,
                "Observer_ID": observer_id,
                "Visit_Date": visit_date,
                "Speech_#": speech["speech_number"],
                "Topic": speech["name"],
                "Token_Position": token_position,
                "Token_Text": match.group(0),
                "Normalized_Token": re.sub(r"[^\w'-]", "", match.group(0)).lower(),
                "Start_Char": token_start,
                "End_Char": token_end,
                "Is_Disfluency": coded_finding is not None,
                "Disfluency_Category": coded_finding["Category"] if coded_finding else "",
                "Coded_Phrase": speech["text"][coded_finding["start"]:coded_finding["end"]] if coded_finding else "",
                "Occurrence_ID": occurrence_ids.get(
                    (coded_finding["start"], coded_finding["end"]), ""
                ) if coded_finding else "",
            })
    return rows


def build_ioa_excel(report_data, speech_results, exclusion_list, n_list, l_list):
    buffer = io.BytesIO()
    coding_rows = build_token_coding_rows(
        report_data[0]["ID"],
        report_data[0]["Observer_ID"],
        report_data[0]["Date"],
        speech_results,
    )
    exclusions = pd.DataFrame({
        "Exclusion_Order": range(1, len(exclusion_list) + 1),
        "Exclusion_Phrase": exclusion_list,
    })
    settings = pd.DataFrame([
        {"Category": "Non-Lexical", "Configured_Item": item} for item in n_list
    ] + [
        {"Category": "Lexical", "Configured_Item": item} for item in l_list
    ])

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        pd.DataFrame(report_data).to_excel(writer, sheet_name="Summary", index=False)
        pd.DataFrame(coding_rows).to_excel(writer, sheet_name="Token_Coding", index=False)
        exclusions.to_excel(writer, sheet_name="Exclusions", index=False)
        settings.to_excel(writer, sheet_name="Coding_Settings", index=False)
    return buffer.getvalue()


# --- IOA CALCULATOR ---
def read_ioa_workbook(uploaded_file):
    sheets = pd.read_excel(uploaded_file, sheet_name=None)
    required_sheets = {"Summary", "Token_Coding", "Exclusions", "Coding_Settings"}
    missing = required_sheets - set(sheets)
    if missing:
        raise ValueError(
            "Missing required sheet(s): " + ", ".join(sorted(missing))
        )

    summary = sheets["Summary"]
    coding = sheets["Token_Coding"]
    required_summary = {"ID", "Observer_ID", "Speech_#", "Topic"}
    required_coding = {
        "Participant_ID", "Observer_ID", "Speech_#", "Topic",
        "Token_Position", "Token_Text", "Normalized_Token",
        "Is_Disfluency", "Disfluency_Category",
    }
    if not required_summary.issubset(summary.columns):
        raise ValueError("The Summary sheet does not match the IOA-ready export format.")
    if not required_coding.issubset(coding.columns):
        raise ValueError("The Token_Coding sheet does not match the IOA-ready export format.")

    return {
        "summary": summary,
        "coding": coding,
        "exclusions": sheets["Exclusions"],
        "settings": sheets["Coding_Settings"],
        "participant": str(summary.iloc[0]["ID"]),
        "observer": str(summary.iloc[0]["Observer_ID"]),
    }


def cell_bool(value):
    if pd.isna(value):
        return False
    if isinstance(value, str):
        return value.strip().lower() in {"true", "yes", "1"}
    return bool(value)


def cell_text(value):
    return "" if pd.isna(value) else str(value)


def align_coding_rows(rows_a, rows_b):
    records_a = rows_a.to_dict("records")
    records_b = rows_b.to_dict("records")
    tokens_a = [cell_text(row["Normalized_Token"]).lower() for row in records_a]
    tokens_b = [cell_text(row["Normalized_Token"]).lower() for row in records_b]
    matcher = SequenceMatcher(None, tokens_a, tokens_b, autojunk=False)
    aligned = []

    for tag, a1, a2, b1, b2 in matcher.get_opcodes():
        if tag == "equal":
            for offset in range(a2 - a1):
                aligned.append((records_a[a1 + offset], records_b[b1 + offset], "match"))
        elif tag == "delete":
            for row_a in records_a[a1:a2]:
                aligned.append((row_a, None, "only_a"))
        elif tag == "insert":
            for row_b in records_b[b1:b2]:
                aligned.append((None, row_b, "only_b"))
        else:
            paired = min(a2 - a1, b2 - b1)
            for offset in range(paired):
                aligned.append((records_a[a1 + offset], records_b[b1 + offset], "different_word"))
            for row_a in records_a[a1 + paired:a2]:
                aligned.append((row_a, None, "only_a"))
            for row_b in records_b[b1 + paired:b2]:
                aligned.append((None, row_b, "only_b"))
    return aligned


def percent(numerator, denominator):
    return round((numerator / denominator) * 100, 2) if denominator else 100.0


def compare_speech(rows_a, rows_b, speech_number, topic, observer_a, observer_b):
    aligned = align_coding_rows(rows_a, rows_b)
    counts = {
        "alignment_positions": len(aligned),
        "matched_words": 0,
        "only_a": 0,
        "only_b": 0,
        "decision_agreements": 0,
        "both_disfluency": 0,
        "either_disfluency": 0,
        "both_coded": 0,
        "category_agreements": 0,
        "disfluency_a": int(rows_a["Is_Disfluency"].apply(cell_bool).sum()),
        "disfluency_b": int(rows_b["Is_Disfluency"].apply(cell_bool).sum()),
    }
    discrepancies = []

    for aligned_position, (row_a, row_b, alignment_type) in enumerate(aligned, start=1):
        if alignment_type != "match":
            if alignment_type == "only_a":
                counts["only_a"] += 1
                difference_type = f"Retained only by {observer_a}"
            elif alignment_type == "only_b":
                counts["only_b"] += 1
                difference_type = f"Retained only by {observer_b}"
            else:
                difference_type = "Different retained word"

            discrepancies.append({
                "Speech_#": speech_number,
                "Topic": topic,
                "Aligned_Position": aligned_position,
                "Difference_Type": difference_type,
                f"{observer_a}_Position": row_a.get("Token_Position", "") if row_a else "",
                f"{observer_a}_Token": row_a.get("Token_Text", "") if row_a else "",
                f"{observer_a}_Disfluency": cell_bool(row_a.get("Is_Disfluency")) if row_a else "",
                f"{observer_a}_Category": cell_text(row_a.get("Disfluency_Category")) if row_a else "",
                f"{observer_b}_Position": row_b.get("Token_Position", "") if row_b else "",
                f"{observer_b}_Token": row_b.get("Token_Text", "") if row_b else "",
                f"{observer_b}_Disfluency": cell_bool(row_b.get("Is_Disfluency")) if row_b else "",
                f"{observer_b}_Category": cell_text(row_b.get("Disfluency_Category")) if row_b else "",
            })
            continue

        counts["matched_words"] += 1
        coded_a = cell_bool(row_a["Is_Disfluency"])
        coded_b = cell_bool(row_b["Is_Disfluency"])
        category_a = cell_text(row_a["Disfluency_Category"])
        category_b = cell_text(row_b["Disfluency_Category"])

        if coded_a == coded_b:
            counts["decision_agreements"] += 1
        if coded_a or coded_b:
            counts["either_disfluency"] += 1
        if coded_a and coded_b:
            counts["both_disfluency"] += 1
            counts["both_coded"] += 1
            if category_a == category_b:
                counts["category_agreements"] += 1

        if coded_a != coded_b or (coded_a and coded_b and category_a != category_b):
            discrepancies.append({
                "Speech_#": speech_number,
                "Topic": topic,
                "Aligned_Position": aligned_position,
                "Difference_Type": "Disfluency coding" if coded_a != coded_b else "Category coding",
                f"{observer_a}_Position": row_a["Token_Position"],
                f"{observer_a}_Token": row_a["Token_Text"],
                f"{observer_a}_Disfluency": coded_a,
                f"{observer_a}_Category": category_a,
                f"{observer_b}_Position": row_b["Token_Position"],
                f"{observer_b}_Token": row_b["Token_Text"],
                f"{observer_b}_Disfluency": coded_b,
                f"{observer_b}_Category": category_b,
            })

    result = {
        "Scope": f"Speech {speech_number}",
        "Topic": topic,
        "Retained_Text_Agreement_%": percent(counts["matched_words"], counts["alignment_positions"]),
        "Disfluency_Decision_Agreement_%": percent(counts["decision_agreements"], counts["matched_words"]),
        "Disfluency_Occurrence_Agreement_%": percent(counts["both_disfluency"], counts["either_disfluency"]),
        "Category_Agreement_%": percent(counts["category_agreements"], counts["both_coded"]),
        "Total_Count_Agreement_%": percent(
            min(counts["disfluency_a"], counts["disfluency_b"]),
            max(counts["disfluency_a"], counts["disfluency_b"]),
        ),
        f"{observer_a}_Retained_Words": len(rows_a),
        f"{observer_b}_Retained_Words": len(rows_b),
        f"{observer_a}_Disfluency_Words": counts["disfluency_a"],
        f"{observer_b}_Disfluency_Words": counts["disfluency_b"],
        "Text_Differences": counts["alignment_positions"] - counts["matched_words"],
        "Coding_Differences": counts["matched_words"] - counts["decision_agreements"],
    }
    return result, discrepancies, counts


def compare_ioa_workbooks(data_a, data_b):
    if data_a["participant"] != data_b["participant"]:
        raise ValueError(
            f"Participant IDs do not match: {data_a['participant']} and {data_b['participant']}."
        )

    speeches_a = set(pd.to_numeric(data_a["coding"]["Speech_#"]).astype(int))
    speeches_b = set(pd.to_numeric(data_b["coding"]["Speech_#"]).astype(int))
    if speeches_a != speeches_b:
        raise ValueError("The files do not contain the same speech numbers.")

    observer_a, observer_b = data_a["observer"], data_b["observer"]
    results, discrepancies = [], []
    for speech_number in sorted(speeches_a):
        rows_a = data_a["coding"][
            pd.to_numeric(data_a["coding"]["Speech_#"]).astype(int) == speech_number
        ].reset_index(drop=True)
        rows_b = data_b["coding"][
            pd.to_numeric(data_b["coding"]["Speech_#"]).astype(int) == speech_number
        ].reset_index(drop=True)
        topic = cell_text(rows_a.iloc[0]["Topic"])
        result, speech_discrepancies, counts = compare_speech(
            rows_a, rows_b, speech_number, topic, observer_a, observer_b
        )
        results.append(result)
        discrepancies.extend(speech_discrepancies)
    return pd.DataFrame(results), pd.DataFrame(discrepancies)


def build_ioa_results_excel(results_df, discrepancies_df, metadata):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        results_df.to_excel(writer, sheet_name="IOA_Results", index=False)
        discrepancies_df.to_excel(writer, sheet_name="Discrepancies", index=False)
        pd.DataFrame([metadata]).to_excel(writer, sheet_name="File_Information", index=False)
    return buffer.getvalue()


def render_ioa_calculator():
    st.header("Interobserver Agreement Calculator")
    st.caption(
        "Upload two IOA-ready coding workbooks. Differences in retained speech "
        "and disfluency coding are evaluated separately."
    )
    col_a, col_b = st.columns(2)
    with col_a:
        file_a = st.file_uploader("Observer A coding file", type=["xlsx"], key="ioa_a")
    with col_b:
        file_b = st.file_uploader("Observer B coding file", type=["xlsx"], key="ioa_b")

    if not file_a or not file_b:
        st.info("Upload both coding files to calculate agreement.")
        return

    try:
        data_a = read_ioa_workbook(file_a)
        data_b = read_ioa_workbook(file_b)
        results_df, discrepancies_df = compare_ioa_workbooks(data_a, data_b)
    except Exception as error:
        st.error(f"Unable to compare these files: {error}")
        return

    observer_a, observer_b = data_a["observer"], data_b["observer"]
    st.markdown(
        f"**Participant:** {data_a['participant']} &nbsp;&nbsp; "
        f"**Observers:** {observer_a} and {observer_b}"
    )
    if observer_a == observer_b:
        st.warning("Both files use the same Observer ID. Confirm that they are independent records.")

    st.subheader("Agreement Results by Speech")
    st.dataframe(results_df, use_container_width=True, hide_index=True)

    st.subheader("Discrepancies")
    if discrepancies_df.empty:
        st.success("No retained-text or coding discrepancies were found.")
    else:
        st.dataframe(discrepancies_df, use_container_width=True, hide_index=True)

    with st.expander("How these measures are calculated"):
        st.markdown(
            "- **Retained-text agreement:** exactly aligned retained words divided by all aligned positions.\n"
            "- **Disfluency decision agreement:** matching disfluency/non-disfluency decisions among words retained by both observers.\n"
            "- **Occurrence agreement:** words coded as disfluent by both observers divided by words coded as disfluent by either observer.\n"
            "- **Category agreement:** matching lexical/non-lexical categories when both observers coded a word as disfluent.\n"
            "- **Total-count agreement:** smaller disfluency-word count divided by the larger count."
        )

    metadata = {
        "Participant_ID": data_a["participant"],
        "Observer_A": observer_a,
        "Observer_B": observer_b,
        "Observer_A_File": file_a.name,
        "Observer_B_File": file_b.name,
    }
    output = build_ioa_results_excel(results_df, discrepancies_df, metadata)
    safe_participant = re.sub(r"[^A-Za-z0-9_-]+", "_", data_a["participant"])
    safe_a = re.sub(r"[^A-Za-z0-9_-]+", "_", observer_a)
    safe_b = re.sub(r"[^A-Za-z0-9_-]+", "_", observer_b)
    st.download_button(
        "📥 Download IOA Results",
        data=output,
        file_name=f"IOA_{safe_participant}_{safe_a}_vs_{safe_b}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# --- 1. SIDEBAR ---
tool_mode = st.sidebar.radio(
    "Select Tool", ["Speech Disfluency Coding", "IOA Calculator"]
)
if tool_mode == "IOA Calculator":
    render_ioa_calculator()
    st.stop()

st.sidebar.header("1. Configure Analysis")
n_input = st.sidebar.text_area("Non-Lexical (N)", value="uh, um, er, ah, mm-hmm, erm, hmm, eh, huh", height=80)
l_input = st.sidebar.text_area("Lexical (L)", value="like, you know, so, therefore, I mean", height=80)

st.sidebar.markdown("---")
st.sidebar.subheader("2. Protocol Exclusions")
st.sidebar.caption("Type any phrase spoken by researchers or protocol instructions to remove them from the analysis (one per line).")
ex_input = st.sidebar.text_area("Words/Phrases to Exclude:", 
                                value="Alright, I'll be starting in 3, 2, 1\nResearcher: Okay, you can begin\nThank you for listening", height=150)

n_list = [w.strip().lower() for w in n_input.split(",")]
l_list = [w.strip().lower() for w in l_input.split(",")]
exclude_list = [p.strip() for p in ex_input.split("\n") if p.strip()]

# --- 2. VISIT METADATA ---
st.header("Visit & Session Information")
c1, c2, c3, c4 = st.columns(4)
with c1: participant_id = st.text_input("Participant ID", value="P001")
with c2: observer_id = st.text_input("Observer ID", value="")
with c3: visit_date = st.date_input("Visit Date", value=datetime.today())
with c4: num_sessions = st.number_input("Number of Speeches", min_value=1, value=1)

sessions_config = []
st.subheader("Speech Timestamps")
cols = st.columns(int(num_sessions))
for i in range(int(num_sessions)):
    with cols[i]:
        st.markdown(f"**Speech {i+1}**")
        s_name = st.text_input(f"Topic", key=f"name_{i}", value=f"Speech {i+1}")
        s_start = st.text_input(f"Start", key=f"start_{i}", placeholder="00:00:00")
        s_end = st.text_input(f"End", key=f"end_{i}", placeholder="00:05:00")
        sessions_config.append({"name": s_name, "start": s_start, "end": s_end})

raw_transcript = st.text_area("3. Paste Transcript Here:", height=200)

# --- 3. ANALYSIS ENGINE ---
def is_filler_heuristic(target, prev, nxt):
    target = target.lower()
    p = re.sub(r'[^\w]', '', prev[-1].lower()) if prev else ""
    n = re.sub(r'[^\w]', '', nxt[0].lower()) if nxt else ""
    if target == "so":
        if p in ["is", "was", "am", "are", "were", "be", "been"]: return False
        functional_so = ["many", "much", "that", "far", "fast", "long", "called"]
        if n in functional_so: return False
        if (n.endswith('y') or n.endswith('ed')) and len(n) > 3: return False
    if target == "like":
        if p in ["i", "you", "he", "she", "it", "we", "they", "to"]: return False
        if n in ["to", "a", "an", "the"]: return False
    return True

def analyze_segment(text_block, n_list, l_list, exclude_list):
    clean_text = clean_transcript_clutter(text_block, exclude_list)
    words = clean_text.split()
    findings = []
    flagged_indices = []
    combined = sorted(n_list + l_list, key=len, reverse=True)

    for target in combined:
        matches = re.finditer(rf'\b{re.escape(target)}\b', clean_text.lower())
        for m in matches:
            if any(i in range(m.start(), m.end()) for i in flagged_indices): continue
            word_idx = len(clean_text[:m.start()].split())
            p_ctx = words[max(0, word_idx-1):word_idx]
            n_ctx = words[word_idx+len(target.split()):word_idx+len(target.split())+1]
            if is_filler_heuristic(target, p_ctx, n_ctx) or target in n_list:
                flagged_indices.extend(range(m.start(), m.end()))
                findings.append({
                    "Word": target, "Category": "Non-Lexical" if target in n_list else "Lexical",
                    "Context": "... " + " ".join(words[max(0, word_idx-2):word_idx+3]) + " ...",
                    "Is Filler?": True, "Count": len(target.split()), "start": m.start(), "end": m.end()
                })
    
    sorted_findings = sorted(findings, key=lambda x: x['start'], reverse=True)
    highlighted = clean_text
    for f in sorted_findings:
        original = highlighted[f['start']:f['end']]
        highlighted = highlighted[:f['start']] + f"**[{original}]**" + highlighted[f['end']:]
    return sorted(findings, key=lambda x: x['start']), words, highlighted

# --- 4. DISPLAY ---
if raw_transcript:
    st.divider()
    report_data = []
    speech_results = []
    lines = raw_transcript.split('\n')
    tabs = st.tabs([s['name'] for s in sessions_config])
    
    for i, tab in enumerate(tabs):
        with tab:
            cfg = sessions_config[i]
            start_s, end_s = get_seconds(cfg['start']), get_seconds(cfg['end'])
            seg_lines = []
            current_time = -1
            for line in lines:
                ts = extract_timestamp(line); 
                if ts is not None: current_time = ts
                if start_s <= current_time <= end_s and not re.match(r'^\d{2}:\d{2}:\d{2}', line.strip()):
                    seg_lines.append(line)
            
            seg_text = " ".join(seg_lines)
            if seg_text.strip():
                findings, words, highlight = analyze_segment(seg_text, n_list, l_list, exclude_list)
                st.subheader("Visual Audit")
                st.markdown(highlight)
                st.divider()
                
                st.subheader(f"Verify Counts: {cfg['name']}")
                df_f = pd.DataFrame(findings) if findings else pd.DataFrame(columns=["Context", "Word", "Category", "Is Filler?"])
                edited_df = st.data_editor(
                    df_f[["Context", "Word", "Category", "Is Filler?"]],
                    key=f"edit_{i}",
                    width=800,
                    disabled=["Context", "Word", "Category"],
                )
                
                confirmed = edited_df[edited_df["Is Filler?"] == True]
                approved_findings = [
                    findings[idx] for idx in confirmed.index
                    if isinstance(idx, int) and 0 <= idx < len(findings)
                ]
                speech_results.append({
                    "speech_number": i + 1,
                    "name": cfg["name"],
                    "text": clean_transcript_clutter(seg_text, exclude_list),
                    "approved_findings": approved_findings,
                })
                total_n = confirmed[confirmed["Category"] == "Non-Lexical"]["Word"].count()
                total_l = sum(confirmed[confirmed["Category"] == "Lexical"]["Word"].apply(lambda x: len(x.split())))
                total_dis = total_n + total_l
                func_words = len(words) - total_dis
                dur_m = (end_s - start_s) / 60
                
                st.markdown("### 📊 Session Summary")
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Dis_per_Min", f"{total_dis/dur_m:.2f}" if dur_m > 0 else "0")
                m2.metric("Dis_per_100_Words", f"{(total_dis/func_words)*100:.2f}" if func_words > 0 else "0")
                m3.metric("Functional Words", func_words)
                m4.metric("Duration (min)", f"{dur_m:.2f}")

                row = {
                    "ID": participant_id,
                    "Observer_ID": observer_id,
                    "Date": visit_date,
                    "Speech_#": i + 1,
                    "Topic": cfg["name"],
                    "Dur": round(dur_m, 2),
                    "Func_Words": func_words,
                    "Dis_per_Min": round(total_dis / dur_m, 2) if dur_m > 0 else 0,
                    "Dis_per_100": round((total_dis / func_words) * 100, 2) if func_words > 0 else 0,
                    "Cleaned_Transcript": clean_transcript_clutter(seg_text, exclude_list),
                }
                for t in (n_list + l_list): row[f"Count_{t}"] = confirmed[confirmed["Word"] == t]["Word"].count()
                report_data.append(row)
            else: st.warning(f"No text found.")

    if report_data:
        st.divider(); st.subheader("4. Final Report Preview")
        final_df = pd.DataFrame(report_data); st.dataframe(final_df)
        if not observer_id.strip():
            st.warning("Enter an Observer ID before downloading the IOA-ready coding file.")
        else:
            ioa_excel = build_ioa_excel(
                report_data, speech_results, exclude_list, n_list, l_list
            )
            safe_participant = re.sub(r"[^A-Za-z0-9_-]+", "_", participant_id)
            safe_observer = re.sub(r"[^A-Za-z0-9_-]+", "_", observer_id)
            st.download_button(
                "📥 Download IOA-Ready Excel Report",
                data=ioa_excel,
                file_name=f"Coding_{safe_participant}_{safe_observer}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        st.markdown("### Reviewed Speech Summary")
        if st.button("Generate Reviewed Speech Word Summary", type="primary"):
            st.session_state["reviewed_speech_docx"] = build_reviewed_speech_docx(
                participant_id, observer_id, visit_date, speech_results
            )

        if "reviewed_speech_docx" in st.session_state:
            safe_id = re.sub(r"[^A-Za-z0-9_-]+", "_", participant_id)
            safe_observer = re.sub(r"[^A-Za-z0-9_-]+", "_", observer_id) or "Observer"
            st.download_button(
                "📄 Download Reviewed Speech (.docx)",
                data=st.session_state["reviewed_speech_docx"],
                file_name=f"Reviewed_Speech_{safe_id}_{safe_observer}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
